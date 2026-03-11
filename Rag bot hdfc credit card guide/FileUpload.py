def fileUpload(file_path, app_name):
    import os

    def extract_text(file_path):
        if "pdf" in file_path:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for i, page in enumerate(reader.pages):
                text += page.extract_text()
            return text    
        elif "docx" in file_path:
            print("Docx File Detected...")
    
            def extract_text(file_path):
                from docx import Document
    
                doc = Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
        elif "txt" in file_path:
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
            return text
        else:
            print("Unsupported file format. Please upload a PDF, DOCX, or TXT file.")
        print("Text File Detected...")

    file_path = os.path.join(os.path.dirname(__file__), file_path)
    text = extract_text(file_path)
    print("Text Extracted...")

    ## Splitter
    from langchain.text_splitter import CharacterTextSplitter
    from langchain.docstore.document import Document

    def split_text_into_chunks(text, chunk_size=500, chunk_overlap=50):
        splitter = CharacterTextSplitter(
            separator="\n", chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        docs = splitter.split_documents([Document(page_content=text)])
        return docs

    chunks = split_text_into_chunks(text)
    print("Text Splitted to chunks...")

    from sentence_transformers import SentenceTransformer

    # Load model (small, fast, and accurate)
    model = SentenceTransformer("all-MiniLM-L6-v2")

    # Get raw text list
    texts = [doc.page_content for doc in chunks]

    # Generate embeddings
    embeddings = model.encode(texts, show_progress_bar=True)
    print("Created Emmbeddings...")
    # store in db
    import faiss
    import numpy as np
    import pickle

    # Convert embeddings to numpy array
    embedding_array = np.array(embeddings).astype("float32")

    # Create FAISS index
    index = faiss.IndexFlatL2(embedding_array.shape[1])  # L2 = Euclidean distance
    index.add(embedding_array)

    # Save FAISS index and texts
    faiss.write_index(index, "faiss_index.index")
    print("Storing Emmbeddings...")
    print("Faiss Index Created`")
    # Save the associated texts for retrieval
    with open(f"{app_name}.pkl", "wb") as f:
        pickle.dump(texts, f)
    print("Saved as .pkl file ")
    return "Model is ready"


fileUpload("HDFC_Credit_Card_Guide.txt",input("Enter the app name: "))
